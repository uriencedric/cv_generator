from docx import Document

from utils import get_current_date_time_yyyymmddhis

# Initialize the Word document
doc = Document()

# Add the header with name and title
doc.add_heading("Chris Makosso", level=1)
doc.add_heading("Sr. Software Engineer", level=2)

# Add location, email, and LinkedIn
doc.add_paragraph("Location: Luxembourg")
doc.add_paragraph("Email: chris.cedric.mak@gmail.com")
doc.add_paragraph("LinkedIn: linkedin.com/in/chris-m-b47800109")

# Add professional summary
doc.add_heading("Professional Summary", level=2)
doc.add_paragraph(
    PROFESSIONAL_SUMMARY_DESC
)

# Add core competencies
doc.add_heading("Core Competencies", level=2)
competencies = COMPETENCIES
for key, values in competencies.items():
    doc.add_paragraph(key, style='List Bullet')
    for value in values:
        doc.add_paragraph(f"- {value}", style='List Bullet 2')

# Add professional experience
doc.add_heading("Professional Experience", level=2)

experiences = EXPERIENCES


for experience in experiences:
    doc.add_heading(f"{experience['role']} | {experience['company']}", level=3)
    doc.add_paragraph(f"{experience['dates']} | {experience['location']}")
    for responsibility in experience['responsibilities']:
        doc.add_paragraph(f"- {responsibility}", style='List Bullet')
    doc.add_paragraph(f"Technologies: {', '.join(experience['technologies'])}")

# Add education
doc.add_heading("Education", level=2)
education = EDUCATION
for edu in education:
    doc.add_paragraph(f"{edu['degree']} | {edu['institution']} | {edu['years']}")

# Add certifications
doc.add_heading("Certifications", level=2)
doc.add_paragraph("Microsoft Certified: Azure Fundamentals")

# Add languages
doc.add_heading("Languages", level=2)
languages = LANGUAGES
for lang, level in languages.items():
    doc.add_paragraph(f"{lang}: {level}")


    # Save the document
file_path = "output/Chris_Makosso_CV_" + get_current_date_time_yyyymmddhis() + ".docx"
doc.save(file_path)

