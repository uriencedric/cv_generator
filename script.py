from docx import Document

from config import *


doc = Document()

for k, v in HEADER.items():
    if k == "name":
        doc.add_heading(v, level=1)
    else:
        doc.add_heading(v, level=2)

for k, v in CONTACT_INFO.items():
    doc.add_paragraph(v)

doc.add_heading("Professional Summary", level=2)
doc.add_paragraph(PROFESSIONAL_SUMMARY_DESC)

doc.add_heading("Core Competencies", level=2)

for key, values in COMPETENCIES.items():
    doc.add_paragraph(key, style='List Bullet')
    for value in values:
        doc.add_paragraph(f" {value}", style='List Bullet 2')

doc.add_heading("Professional Experience", level=2)

for experience in EXPERIENCES:
    doc.add_heading(f"{experience['role']} | {experience['company']}", level=3)
    doc.add_paragraph(f"{experience['dates']} | {experience['location']}")
    for responsibility in experience['responsibilities']:
        doc.add_paragraph(f" {responsibility}", style='List Bullet')
    doc.add_paragraph(f"Technologies: {', '.join(experience['technologies'])}")

doc.add_heading("Education", level=2)

for edu in EDUCATION:
    doc.add_paragraph(f"{edu['degree']} | {edu['institution']} | {edu['years']}")

doc.add_heading("Certifications", level=2)
doc.add_paragraph(CERTIFICATION_PARAGRAPH)

doc.add_heading("Languages", level=2)
languages = LANGUAGES
for lang, level in languages.items():
    doc.add_paragraph(f"{lang}: {level}")

doc.save(OUTPUT_PATH)
